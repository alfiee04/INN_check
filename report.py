

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import config






def _set_font(run, size_pt: int, bold: bool = False, color: RGBColor | None = None):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def _heading(doc: Document, text: str, level: int, size_pt: int,
             align=WD_ALIGN_PARAGRAPH.LEFT, bold: bool = True):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    _set_font(run, size_pt, bold=bold)
    return p


def _add_horizontal_rule(doc: Document):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)






def _build_title_page(doc: Document, company_name: str, inn: str,
                      executor_name: str, address: str):
    doc.add_paragraph()
    doc.add_paragraph()

    _heading(doc, "СКРИНШОТЫ", level=1, size_pt=28,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    _heading(doc, "проверяемого контрагента", level=2, size_pt=18,
             align=WD_ALIGN_PARAGRAPH.CENTER, bold=False)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p.add_run(company_name), 18, bold=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p2.add_run(f"ИНН {inn}"), 14)

    doc.add_paragraph()
    doc.add_paragraph()


    tbl = doc.add_table(rows=3, cols=2)
    tbl.style = "Table Grid"
    labels = ["Исполнитель", "Дата проверки", "Юридический адрес"]
    values = [
        executor_name,
        datetime.now().strftime("%d.%m.%Y %H:%M"),
        address,
    ]
    for i, (lbl, val) in enumerate(zip(labels, values)):
        row = tbl.rows[i]
        _set_cell_bg(row.cells[0], "D9E2F3")
        lbl_run = row.cells[0].paragraphs[0].add_run(lbl)
        _set_font(lbl_run, 11, bold=True)
        val_run = row.cells[1].paragraphs[0].add_run(val)
        _set_font(val_run, 11)

    doc.add_page_break()






def _build_address_section(doc: Document, address: str):
    _heading(doc, "1. Подтверждение адреса местонахождения организации",
             level=2, size_pt=14)
    p = doc.add_paragraph()
    _set_font(p.add_run(f"Юридический и почтовый адрес: {address}"), 12)
    _add_horizontal_rule(doc)
    doc.add_paragraph()






def _build_screenshot_section(doc: Document, idx: int, title: str,
                               screenshot_path: str | None, error_msg: str | None):
    _heading(doc, f"{idx}. {title}", level=2, size_pt=13)

    now_str = datetime.now().strftime("%d.%m.%Y %H:%M:%S")

    if screenshot_path and os.path.exists(screenshot_path):
        try:
            doc.add_picture(screenshot_path, width=Inches(6.3))
            p = doc.add_paragraph()
            _set_font(
                p.add_run(f"Скриншот сделан: {now_str}"),
                10, bold=False, color=RGBColor(0x55, 0x55, 0x55)
            )
        except Exception as e:
            p = doc.add_paragraph()
            _set_font(p.add_run(f"[Ошибка вставки скриншота: {e}]"), 11)
    else:
        p = doc.add_paragraph()
        if error_msg and str(error_msg).startswith("MANUAL:"):
            run = p.add_run("⚠️  Раздел предусмотрен ТЗ и заполняется вручную")
            _set_font(run, 12, bold=True, color=RGBColor(0x77, 0x77, 0x77))
            p2 = doc.add_paragraph()
            _set_font(p2.add_run(str(error_msg).replace("MANUAL:", "").strip()), 10,
                      color=RGBColor(0x55, 0x55, 0x55))
        else:
            run = p.add_run("❌  Сайт временно недоступен или не загрузился")
            _set_font(run, 12, bold=True, color=RGBColor(0xC0, 0x00, 0x00))
            if error_msg:
                p2 = doc.add_paragraph()
                _set_font(p2.add_run(f"Ошибка: {error_msg[:300]}"), 10,
                          color=RGBColor(0x88, 0x00, 0x00))

    _add_horizontal_rule(doc)
    doc.add_paragraph()






def _build_footer(doc: Document):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ts = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
    _set_font(
        p.add_run(f"Отчёт сформирован автоматически · {ts}"),
        9, color=RGBColor(0x88, 0x88, 0x88)
    )






def create_word_report(
    company_name: str,
    inn: str,
    executor_name: str,
    address: str,
    screenshots_data: list[tuple[str, str | None, str | None]],
) -> str:
    doc = Document()


    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    _build_title_page(doc, company_name, inn, executor_name, address)
    _build_address_section(doc, address)

    for idx, (title, path, err) in enumerate(screenshots_data, start=2):
        _build_screenshot_section(doc, idx, title, path, err)

    _build_footer(doc)

    filename = f"report_{inn}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)
    return filename
